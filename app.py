from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from neo4j import GraphDatabase
from pydantic import BaseModel
import os
import pdfplumber
import docx
import uvicorn
from fastapi.middleware.cors import CORSMiddleware
import re
from neo4j import GraphDatabase
from typing import List
from typing_extensions import TypedDict
import ollama
from langchain_groq import ChatGroq
from langchain_core.documents import Document
from langchain_neo4j import Neo4jGraph, GraphCypherQAChain
from langchain_experimental.graph_transformers import LLMGraphTransformer
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv

load_dotenv()

NEO4J_URI = os.getenv("NEO4J_URI")
NEO4J_USER = os.getenv("NEO4J_USERNAME")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
graph = Neo4jGraph(refresh_schema=True)


class ChatRequest(BaseModel):
    query: str


def construct_knowledge_graph_from_text(file_path: str):
    """Processes text file to construct and populate a knowledge graph."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    llm = ChatOpenAI(temperature=0, model_name="gpt-4-turbo")
    llm_transformer = LLMGraphTransformer(llm=llm)
    documents = [Document(page_content=text)]
    graph_documents = llm_transformer.convert_to_graph_documents(documents)

    graph.add_graph_documents(
        graph_documents, baseEntityLabel=True, include_source=True
    )
    print("Knowledge graph constructed and populated successfully.")


def normalize_name(name: str) -> str:
    """Normalize names by replacing spaces with underscores."""
    return re.sub(r"[\s]+", "_", name.strip()).lower()


def denormalize_name(name: str, original_name: str, schema: list) -> str:
    """Reverse normalization: replace underscores with spaces if the original name had spaces."""
    # If the name is an entity or relationship, we do not denormalize it
    if name in schema:
        return name
    return original_name.replace("_", " ")


def replace_names_in_query(
    query: str, schema_entities: list, schema_relationships: list
) -> str:
    """
    Replace entity and relationship names in the query with their correct forms from the schema.
    Searches for entities and relationships in the query text and replaces them if found.
    Preserves the capitalization of Cypher keywords.
    """
    # Create a mapping of normalized entity names to original names
    entity_mapping = {normalize_name(entity): entity for entity in schema_entities}
    relationship_mapping = {
        normalize_name(relationship): relationship
        for relationship in schema_relationships
    }

    # Function to replace matching text (entity or relationship) while preserving case of the rest of the query
    def replace_match(match):
        matched_text = match.group(0)
        normalized_match = normalize_name(matched_text)

        # If the match is an entity or relationship, replace it with the correct form
        if normalized_match in entity_mapping:
            return denormalize_name(
                entity_mapping[normalized_match], matched_text, schema_entities
            )
        elif normalized_match in relationship_mapping:
            return denormalize_name(
                relationship_mapping[normalized_match],
                matched_text,
                schema_relationships,
            )

        return matched_text

    # Use regular expression to replace entities and relationships without affecting other parts of the query
    transformed_query = re.sub(r"\b[\w\s-]+\b", replace_match, query)
    return transformed_query


def get_graph_schema(uri, user, password):
    """Retrieve labels, relationships, properties, and IDs from Neo4j."""
    driver = GraphDatabase.driver(uri, auth=(user, password))
    schema = {
        "entities": [],
        "relationships": [],
        "node_properties": {},
        "relationship_properties": {},
        "entity_ids": {},
    }

    with driver.session() as session:
        # Get all labels (entities)
        labels_query = "CALL db.labels() YIELD label RETURN label"
        schema["entities"] = [record["label"] for record in session.run(labels_query)]

        # Get all relationship types
        relationships_query = (
            "CALL db.relationshipTypes() YIELD relationshipType RETURN relationshipType"
        )
        schema["relationships"] = [
            record["relationshipType"] for record in session.run(relationships_query)
        ]

        # Get node properties and IDs for each label
        for label in schema["entities"]:
            # Retrieve distinct properties
            properties_query = f"""
            MATCH (n:`{label}`)
            UNWIND keys(n) AS property
            RETURN DISTINCT property
            """
            properties = [
                record["property"] for record in session.run(properties_query)
            ]
            schema["node_properties"][label] = properties

            # Retrieve IDs of nodes with this label
            ids_query = f"""
            MATCH (n:`{label}`)
            RETURN id(n) AS node_id, n.id AS entity_id LIMIT 100
            """
            ids = [
                {"node_id": record["node_id"], "entity_id": record.get("entity_id")}
                for record in session.run(ids_query)
            ]
            schema["entity_ids"][label] = ids

        # Get relationship properties for each relationship type
        for rel_type in schema["relationships"]:
            properties_query = f"""
            MATCH ()-[r:`{rel_type}`]->()
            UNWIND keys(r) AS property
            RETURN DISTINCT property
            """
            properties = [
                record["property"] for record in session.run(properties_query)
            ]
            schema["relationship_properties"][rel_type] = properties

    driver.close()
    return schema


# def query_knowledge_graph_with_schema(question: str, schema: dict):
#     """Generate a Cypher query using Ollama with schema and entity IDs."""
#     # Prepare schema information
#     entities = ", ".join(schema["entities"])
#     relationships = ", ".join(schema["relationships"])
#     entity_id_mapping = {
#         label: [entry["entity_id"] for entry in schema["entity_ids"].get(label, [])]
#         for label in schema["entities"]
#     }

#     # Schema Description for Ollama
#     schema_description = f"""
#     The graph schema includes:
#     - Entities: {entities}.
#     - Relationships: {relationships}.
#     Entity IDs for each label:
#     {', '.join([f'{label}: {ids}' for label, ids in entity_id_mapping.items()])}
#     Use these entities, relationships, and IDs to generate the query.
#     """

#     # Ollama Prompt
#     prompt = f"""
#         Generate a valid Cypher query for Neo4j to answer the question:
#         "{question}"
#         Schema details:
#         {schema_description}
#         Use the property `entity_id` to match specific entities. For example:
#         MATCH (n {{entity_id: "Artificial_Intelligence"}}) RETURN n
#         Only output the Cypher query, no explanation or comments are required.
#     """

#     response = ollama.generate(prompt=prompt, model="llama3.1")
#     print("Response from server:", response["response"])
#     cypher_query = response["response"].strip()

#     # Replace entity names with their IDs
#     # cypher_query = replace_names_with_ids(cypher_query, entity_id_mapping)
#     # print("Generated Cypher Query:", cypher_query)
#     return cypher_query


def query_knowledge_graph_with_schema(question: str, schema: dict):
    """Generate a Cypher query using Ollama with schema and entity IDs."""
    # Prepare schema information
    entities = ", ".join(schema["entities"])
    relationships = ", ".join(schema["relationships"])

    # Schema Description for Ollama
    schema_description = f"""
    The graph schema includes:
    - Entities: {entities}.
    - Relationships: {relationships}.
    Use the property `id` to match specific entities. For example:
    MATCH (n {{id: "Artificial_Intelligence"}}) RETURN n
    Only output the Cypher query.
    """

    # Ollama Prompt
    prompt = f"""
        Generate a valid Cypher query for Neo4j to answer the question:
        "{question}"
        Schema details:
        {schema_description}
        Use the property `id` for entity matching. For example:
        MATCH (n {{id: "Artificial_Intelligence"}}) RETURN n
        Only output the Cypher query, no explanation or comments.
    """

    response = ollama.generate(prompt=prompt, model="llama3.1")
    print("Response from server:", response["response"])
    cypher_query = response["response"].strip()
    cypher_query = response["response"].replace("`", "")

    # Ensure the query uses single curly braces for valid Cypher syntax
    return cypher_query


def replace_names_with_ids(cypher_query: str, entity_id_mapping: dict) -> str:
    """Replace entity names with their entity_id in the Cypher query."""
    for label, ids in entity_id_mapping.items():
        for entity_id in ids:
            if entity_id in cypher_query:
                cypher_query = cypher_query.replace(
                    entity_id, f'match (n {{entity_id: "{entity_id}"}})'
                )
    return cypher_query


def execute_cypher_query(uri, user, password, cypher_query):
    """Execute the Cypher query against Neo4j."""
    driver = GraphDatabase.driver(uri, auth=(user, password))
    result = []

    with driver.session() as session:
        try:
            query_result = session.run(cypher_query)

            result = [record.data() for record in query_result]
            print(result)
        except Exception as e:
            print(f"Error: {str(e)}\nQuery: {cypher_query}")
            result = []

    driver.close()
    return result


@app.post("/chatbot")
async def chatbot(request: ChatRequest):
    user_question = request.query

    try:
        # Step 1: Retrieve the graph schema
        schema = get_graph_schema(NEO4J_URI, NEO4J_USER, NEO4J_PASSWORD)

        # Step 2: Generate Cypher query using schema and question
        cypher_query = query_knowledge_graph_with_schema(user_question, schema)

        # Ensure the query uses valid syntax
        cypher_query = cypher_query.replace("entity_id", "id")
        print("Generated Cypher Query:", cypher_query)

        # Step 3: Execute the Cypher query
        results = execute_cypher_query(
            NEO4J_URI, NEO4J_USER, NEO4J_PASSWORD, cypher_query
        )

        if not results:
            return {"response": "No results found for your query."}

        # Step 4: Format results into a context for Ollama
        formatted_results = "\n".join([str(result) for result in results])
        prompt = f"""
            Based on the following query results, generate a concise and structured answer:

            Query Results:
            {formatted_results}

            User Question:
            "{user_question}"

            Provide the answer in a structured format.
            Do not include your own knowledge just use the answer from query results and format it in a readable format.
        """

        # Step 5: Pass the results to Ollama
        response = ollama.generate(prompt=prompt, model="llama3.1")
        structured_answer = response["response"].strip()
        print("Structured Answer:", structured_answer)

        return {"response": structured_answer}

    except Exception as e:
        print(f"Error: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"response": f"An error occurred: {str(e)}"},
        )


@app.get("/analytics")
async def get_analytics():
    with driver.session() as session:
        total_nodes_result = session.run("MATCH (n) RETURN count(n) AS total_nodes")
        total_nodes = total_nodes_result.single()["total_nodes"]

        total_relationships_result = session.run(
            "MATCH ()-[r]->() RETURN count(r) AS total_relationships"
        )
        total_relationships = total_relationships_result.single()["total_relationships"]

        if total_nodes > 1:
            graph_density = total_relationships / (total_nodes * (total_nodes - 1))
        else:
            graph_density = 0.0

        if total_nodes > 0:
            avg_degree = total_relationships / total_nodes
        else:
            avg_degree = 0.0

        largest_component_result = session.run(
            """
            MATCH (n)
            WITH n, id(n) AS nodeId
            MATCH (n)-[*]-(m)
            WITH nodeId, count(DISTINCT m) AS component_size
            ORDER BY component_size DESC LIMIT 1
            RETURN component_size
            """
        )
        largest_component_size = largest_component_result.single()["component_size"]

    return {
        "totalNodes": total_nodes,
        "totalRelationships": total_relationships,
        "graphDensity": graph_density,
        "avgDegree": avg_degree,
        "largestComponentSize": largest_component_size,
    }


def extract_text(file_path: str, file_type: str) -> str:
    extracted_text = ""
    if file_type == "text/plain":
        with open(file_path, "r", encoding="utf-8") as file:
            extracted_text = file.read()
    elif file_type == "application/pdf":
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted_text += page.extract_text() + "\n"
    elif file_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            extracted_text += paragraph.text + "\n"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format.")
    return extracted_text.strip()


def preprocess_text(text):
    text = text.lower()
    text = re.sub(r"[^\w\s.,!?;]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Save the uploaded file
        file_location = f"uploads/{file.filename}"
        os.makedirs("uploads", exist_ok=True)
        with open(file_location, "wb") as buffer:
            buffer.write(await file.read())

        # Step 1: Extract text from the uploaded file
        extracted_text = extract_text(file_location, file.content_type)
        print("Extracted Text:")
        print(extracted_text)

        # Step 2: Preprocess the text
        preprocessed_text = preprocess_text(extracted_text)
        print("Preprocessed Text:")
        print(preprocessed_text)

        # Step 3: Construct the knowledge graph from the preprocessed text
        temp_text_file = "uploads/temp_text_file.txt"
        with open(temp_text_file, "w", encoding="utf-8") as temp_file:
            temp_file.write(preprocessed_text)

        construct_knowledge_graph_from_text(temp_text_file)

        # Optional: Clean up the temporary files
        os.remove(temp_text_file)
        os.remove(file_location)

        return {
            "message": "File processed and knowledge graph constructed successfully.",
            "extracted_text": preprocessed_text,
        }

    except Exception as e:
        return JSONResponse(status_code=500, content={"message": str(e)})


@app.post("/run-query")
async def run_query(query: ChatRequest):
    cypher_query = query.query
    try:
        with driver.session() as session:
            query_result = session.run(cypher_query)
            result = [record.data() for record in query_result]
            return {"success": True, "result": result}
    except Exception as e:
        return JSONResponse(
            status_code=400, content={"success": False, "error": str(e)}
        )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
